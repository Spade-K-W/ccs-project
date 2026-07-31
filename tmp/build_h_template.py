from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


OUT = Path(r"C:\Users\mateb\Desktop\ccs_project\ccs-project\output\docx\H题_车载平衡滚球运动控制系统_设计报告模板.docx")

BLACK = "000000"
GRAY = "666666"
LIGHT_GRAY = "E7E6E6"
PALE_GRAY = "F4F4F4"
USABLE_DXA = 8957  # 15.8 cm at A4 with the margins below


def set_cell_margins(cell, top=72, start=90, bottom=72, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in("w:tcW")
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")

    tblGrid = table._tbl.tblGrid
    for child in list(tblGrid):
        tblGrid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tblGrid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge)
        tag = qn(f"w:{edge}")
        element = tcBorders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_run_font(run, east_asia="宋体", size=10.5, bold=False, color=BLACK, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), east_asia)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")


def set_style_font(style, east_asia, size, bold=False, color=BLACK):
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), east_asia)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, "宋体", 10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    for name, size, before, after in (
        ("Heading 1", 14, 12, 6),
        ("Heading 2", 12, 9, 4),
        ("Heading 3", 10.5, 6, 3),
    ):
        style = doc.styles[name]
        set_style_font(style, "黑体", size, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing = 1.15

    caption = doc.styles["Caption"]
    set_style_font(caption, "宋体", 9)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.first_line_indent = Cm(0)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        set_style_font(style, "宋体", 10.5)
        style.paragraph_format.left_indent = Cm(0.74)
        style.paragraph_format.first_line_indent = Cm(-0.37)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.2

    if "占位提示" not in doc.styles:
        style = doc.styles.add_style("占位提示", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["占位提示"]
    set_style_font(style, "宋体", 9.5, color=GRAY)
    style.paragraph_format.first_line_indent = Cm(0)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)

    if "公式" not in doc.styles:
        style = doc.styles.add_style("公式", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["公式"]
    set_style_font(style, "宋体", 10.5)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.first_line_indent = Cm(0)
    style.paragraph_format.space_before = Pt(3)
    style.paragraph_format.space_after = Pt(3)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=9)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_para(doc, text="", style=None, bold_prefix=None, align=None, first_indent=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if first_indent is not None:
        p.paragraph_format.first_line_indent = first_indent
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r, east_asia="宋体", size=9.5 if style == "占位提示" else 10.5,
                     color=GRAY if style == "占位提示" else BLACK,
                     italic=(style == "占位提示"))
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, east_asia="黑体", size={1: 14, 2: 12, 3: 10.5}[level], bold=True)
    return p


def add_prompt(doc, text):
    return add_para(doc, f"【填写提示】{text}", style="占位提示")


def add_formula(doc, text):
    p = doc.add_paragraph(style="公式")
    r = p.add_run(text)
    set_run_font(r, east_asia="宋体", size=10.5)
    return p


def add_figure_placeholder(doc, label, prompt, height_lines=3):
    outer = doc.add_table(rows=1, cols=1)
    set_table_width(outer, [USABLE_DXA])
    outer_cell = outer.cell(0, 0)
    set_cell_margins(outer_cell, top=0, start=0, bottom=0, end=0)
    no_border = {"val": "nil", "sz": "0", "space": "0", "color": "FFFFFF"}
    set_cell_border(outer_cell, top=no_border, bottom=no_border, left=no_border, right=no_border)
    spacer = outer_cell.paragraphs[0]
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    spacer.paragraph_format.line_spacing = Pt(1)

    table = outer_cell.add_table(rows=1, cols=1)
    set_table_width(table, [USABLE_DXA])
    cell = table.cell(0, 0)
    shade_cell(cell, "FAFAFA")
    border = {"val": "dashed", "sz": "6", "space": "0", "color": "A6A6A6"}
    set_cell_border(cell, top=border, bottom=border, left=border, right=border)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"［插入{label}］\n{prompt}")
    set_run_font(r, east_asia="宋体", size=9.5, color=GRAY)
    for _ in range(max(0, height_lines - 2)):
        p.add_run("\n")
    cap = outer_cell.add_paragraph(style="Caption")
    rr = cap.add_run(label)
    set_run_font(rr, east_asia="宋体", size=9)
    return outer


def add_table(doc, headers, rows, widths, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        set_run_font(r, east_asia="黑体", size=font_size, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, east_asia="宋体", size=font_size,
                         color=GRAY if "填写" in text or "替换" in text else BLACK,
                         italic=("填写" in text or "替换" in text))
    # Keep compact technical tables together when they fit on a page.
    for row in table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
    return table


def add_note_box(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [USABLE_DXA])
    cell = table.cell(0, 0)
    shade_cell(cell, PALE_GRAY)
    border = {"val": "single", "sz": "4", "space": "0", "color": "BFBFBF"}
    set_cell_border(cell, top=border, bottom=border, left=border, right=border)
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(title)
    set_run_font(r1, east_asia="黑体", size=9.5, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, east_asia="宋体", size=9.5, color=GRAY)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r, east_asia="宋体", size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r, east_asia="宋体", size=10.5)
    return p


def add_test_table(doc, headers, row_count, widths, requirement, note=None):
    rows = []
    for i in range(1, row_count + 1):
        rows.append([str(i)] + ["［填写］"] * (len(headers) - 2) + [requirement])
    add_table(doc, headers, rows, widths, font_size=8.5)
    if note:
        add_prompt(doc, note)


doc = Document()
configure_styles(doc)

section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Cm(2.3)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.6)
section.right_margin = Cm(2.6)
section.header_distance = Cm(1.0)
section.footer_distance = Cm(1.2)
section.different_first_page_header_footer = True

footer = section.footer
add_page_field(footer.paragraphs[0])
section.first_page_footer.paragraphs[0].text = ""

props = doc.core_properties
props.title = "H题 车载平衡滚球运动控制系统——设计报告模板"
props.subject = "全国大学生电子设计竞赛设计报告模板"
props.author = "参赛队伍"
props.keywords = "循迹小车；滚球控制；视觉检测；PID；电子设计竞赛"

# Cover / abstract
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("2026年全国大学生电子设计竞赛")
set_run_font(r, east_asia="黑体", size=15, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("H．车载平衡滚球运动控制系统")
set_run_font(r, east_asia="黑体", size=22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
r = p.add_run("设 计 报 告")
set_run_font(r, east_asia="黑体", size=18, bold=True)

meta_rows = [
    ["参赛学校", "［填写］", "队号", "［填写］"],
    ["参赛队员", "［填写］", "完成日期", "［填写］"],
]
meta = add_table(doc, ["项目", "内容", "项目", "内容"], meta_rows, [1500, 3000, 1500, 2957], font_size=9)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(5)
r = p.add_run("摘  要")
set_run_font(r, east_asia="黑体", size=14, bold=True)

abstract_text = (
    "本系统以［主控芯片型号］为核心，设计了一辆搭载平衡滚球运动控制装置的轮式循线小车。"
    "小车采用八路红外光电传感器获取黑线位置，结合编码器与惯性测量单元完成直线、半圆弧及启停状态判定；"
    "滚球装置利用摄像头提取钢球坐标，通过［舵机/电机/其他机构］调节摆杆倾角，实现钢球定点、往返及车辆运动过程中的稳定控制。"
    "系统同时完成按键启动、时间显示以及车载图传记录。测试结果表明：［填写关键指标，例如整圈时间、停车偏差、最大滚球误差］，"
    "满足题目主要指标要求。"
)
add_para(doc, abstract_text)
add_prompt(doc, "摘要建议控制在200～300字，依次写“系统构成—关键方法—测试结果”；把方括号内容替换为实物和实测值。")

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.space_before = Pt(5)
r1 = p.add_run("关键词：")
set_run_font(r1, east_asia="黑体", size=10.5, bold=True)
r2 = p.add_run("循迹小车；平衡滚球；视觉检测；PID控制；图像传输")
set_run_font(r2, east_asia="宋体", size=10.5)

add_note_box(
    doc,
    "模板说明：",
    "灰色“填写提示”和方括号内容仅供撰写时参考，定稿前应全部替换或删除。当前工程中可确认的器件已作为建议项出现，实际型号以最终实物为准。"
)

doc.add_page_break()

# Chapter 1
add_heading(doc, "1 设计任务及要求", 1)
add_heading(doc, "1.1 任务分析", 2)
add_para(
    doc,
    "本题要求在尺寸不超过35 cm×25 cm的轮式小车上安装带凹槽摆杆、角度控制机构、钢球及图传装置。"
    "小车仅使用红外光电模块进行黑线循迹，在由两段1.5 m直线与两段半径0.5 m半圆组成的环形线路上完成定时、定点和整圈运行；"
    "同时通过摄像头检测钢球在25 cm摆杆中的位置，并在车辆静止或行驶时将钢球稳定在指定坐标附近。"
)
add_para(
    doc,
    "系统可分为小车运动控制、滚球位置控制、视觉检测与图传、计时显示四个子系统。报告应围绕“位置如何检测、状态如何切换、"
    "控制量如何计算、指标如何验证”展开，并使每一项测试结果能够对应题目评分点。"
)

add_heading(doc, "1.2 主要技术指标与验证材料", 2)
requirements = [
    ["1", "图传实时显示、完整录像并可回放", "画面覆盖全摆杆；录像文件与回放截图"],
    ["2", "顺时针一圈并停在A点", "总时间≤20 s；停车偏差≤2 cm"],
    ["3", "静止时O→+5 cm→−5 cm并稳定", "总时间≤5 s；±5 cm最大误差≤1 cm"],
    ["4", "A到B行驶并稳定在O点", "AB时间≤8 s；最大位置误差≤1 cm"],
    ["5", "整圈行驶并稳定在O点", "整圈≤30 s；最大位置误差≤1 cm"],
    ["6", "整圈行驶并稳定在任意设定位置", "整圈≤30 s；最大位置误差≤1 cm"],
    ["其他", "启动按键、≤2英寸显示、车载供电、尺寸约束", "实物照片、尺寸与功耗记录"],
]
add_table(doc, ["序号", "任务内容", "报告中应给出的证据"], requirements, [800, 3550, 4607], font_size=8.7)
add_prompt(doc, "本节只概括要求，不要重复粘贴整份题目；尺寸、路线和关键限制应明确写出。")

add_heading(doc, "1.3 设计目标与难点", 2)
for text in (
    "直线与半圆弧的循迹参数和基础速度不同，需要设计可靠的路段状态机，避免误触发和丢线。",
    "车体加减速与转弯产生的惯性扰动会使钢球滚动，滚球控制必须兼顾响应速度、超调和抗扰性。",
    "摄像头既是滚球位置检测的唯一传感器，又要满足图传录像要求，需处理延迟、标定和光照变化。",
    "各测试需独立计时并保存钢球位置曲线，便于计算最大误差并形成可复核的数据。",
):
    add_bullet(doc, text)

# Chapter 2
add_heading(doc, "2 系统方案设计", 1)
add_heading(doc, "2.1 系统总体框架", 2)
add_para(
    doc,
    "系统由主控制器、红外循迹模块、编码器与MPU6050、四轮电机驱动、摆杆执行机构、视觉处理/图传模块、"
    "按键及显示模块组成。主控制器周期性读取循迹误差和车辆姿态，输出左右轮目标速度；视觉模块输出钢球位置，"
    "滚球控制器根据设定点与实测位置调节摆杆倾角。"
)
add_figure_placeholder(
    doc,
    "图1 系统总体框图",
    "建议画出：红外/编码器/MPU6050/摄像头 → 主控或视觉处理器 → 电机驱动/摆杆执行机构/OLED与数码管/图传接收端。",
    height_lines=4,
)
add_prompt(doc, "图中必须区分“循迹位置检测”和“钢球位置检测”：前者为红外模块，后者按题目要求必须为摄像头。")

add_heading(doc, "2.2 主控制器方案", 2)
add_para(
    doc,
    "主控制器需要完成多路GPIO采样、电机PWM、编码器测速、惯性传感器通信、按键与显示刷新，并与视觉处理模块交换钢球坐标。"
    "本系统选用［MSPM0G3507/实际型号］，其定时器、ADC/通信接口和运算能力可满足控制周期要求。"
)
add_table(
    doc,
    ["候选方案", "优点", "不足", "结论"],
    [
        ["方案一：［型号］", "［填写］", "［填写］", "［选用/不选］"],
        ["方案二：MSPM0G3507", "接口资源较丰富，适合多定时任务", "需要完成外设驱动与调度", "［按实物确认］"],
    ],
    [1700, 2600, 2600, 2057],
    font_size=8.5,
)

add_heading(doc, "2.3 小车驱动与测速方案", 2)
add_para(
    doc,
    "小车采用［两轮差速/四轮差速］结构，电机驱动器为［TB6612FNG/实际型号］。编码器测得各轮转速，速度闭环根据目标速度修正PWM。"
    "为降低启动冲击，可采用斜坡给定；停止阶段根据对应任务选择直接制动或速度斜坡下降。"
)
add_figure_placeholder(doc, "图2 小车驱动与编码器连接框图", "标注电池、稳压、电机驱动、四个电机、编码器与主控连接。")

add_heading(doc, "2.4 红外循迹方案", 2)
add_para(
    doc,
    "循迹模块采用八路红外光电传感器，沿车头横向排列。各通道根据黑/白电平转换为二值量，再通过加权求和得到横向偏差。"
    "直线段与弧线段分别设置基础速度及控制参数，路段切换由红外特征、编码器距离和航向角共同约束。"
)
add_table(
    doc,
    ["方案", "特点", "适用性"],
    [
        ["单通道/少量传感器", "电路简单，但偏差分辨率低", "不利于高速和弧线连续控制"],
        ["八路红外加权", "可获得离散横向误差，运算量小", "符合题目限制，作为本系统方案"],
    ],
    [2200, 3600, 3157],
    font_size=8.7,
)

add_heading(doc, "2.5 惯性测量与路段判断方案", 2)
add_para(
    doc,
    "MPU6050用于估计累计航向变化，编码器用于估算行驶距离与速度。二者不替代红外循迹，而是为直线/弧线状态切换、"
    "整圈结束以及异常保护提供附加判据。应在报告中写明零偏校准时间、角度积分周期、距离换算系数及阈值来源。"
)

add_heading(doc, "2.6 平衡滚球装置方案", 2)
add_heading(doc, "2.6.1 摆杆与角度执行机构", 3)
add_para(
    doc,
    "摆杆由长度25 cm的4分PPR水管加工而成，左端铰接且离车板高度h≥5 cm，右端由［舵机连杆/丝杆电机/其他机构］改变高度或角度。"
    "执行机构需覆盖滚球所需倾角范围，并尽量减小回差。"
)
add_figure_placeholder(doc, "图3 摆杆机械结构及尺寸", "建议使用实物侧视图或CAD图，标注铰点、h、杆长、执行机构和摄像头位置。")

add_heading(doc, "2.6.2 钢球视觉检测与标定", 3)
add_para(
    doc,
    "摄像头固定在凹槽上方，画面覆盖整个摆杆。视觉处理流程可包括裁剪感兴趣区域、颜色/灰度分割、形态学处理、轮廓筛选与圆心计算，"
    "最后通过像素—长度标定得到钢球相对O点的坐标。"
)
add_figure_placeholder(doc, "图4 钢球视觉识别流程与标定画面", "放置原图、二值图、圆心叠加图以及−5 cm、0、+5 cm标定点。")

add_heading(doc, "2.6.3 图传与录像方案", 3)
add_para(
    doc,
    "图传发送端和摄像头稳固安装在车体，接收端连接［笔记本/PAD/其他设备］。测试时实时显示全摆杆画面，"
    "按“日期—任务—测试序号”保存视频，并在报告中给出帧率、分辨率、可用传输距离和录像文件示例。"
)

add_heading(doc, "2.7 人机交互与供电方案", 2)
add_para(
    doc,
    "系统设置启动按键与不大于2英寸的显示装置。OLED用于显示模式、状态或触发原因，数码管用于显示从按键启动到车辆停止的时间。"
    "车载电池经［填写稳压方案］分别为电机、主控、视觉及执行机构供电，必要时采用分区供电和共地以降低电机干扰。"
)

# Chapter 3
add_heading(doc, "3 理论分析与控制方法", 1)
add_heading(doc, "3.1 小车循迹控制理论", 2)
add_heading(doc, "3.1.1 八路红外加权误差", 3)
add_para(
    doc,
    "设第i路传感器检测黑线时bᵢ=1，否则bᵢ=0；其横向权值为wᵢ。中间通道权值接近0，左、右通道分别取负值和正值。"
    "当至少一路检测到黑线时，离散横向误差可表示为："
)
add_formula(doc, "e(k) = [Σ bᵢ(k)·wᵢ] / [Σ bᵢ(k)]")
add_para(
    doc,
    "权值应结合传感器从左到右的实际编号填写。若某通道长期异常，应先排查硬件；软件屏蔽需在报告中说明，并重新验证误差的单调性。"
)
add_table(
    doc,
    ["通道", "X1", "X2", "X3", "X4", "X5", "X6", "X7", "X8"],
    [["权值", "［填］", "［填］", "［填］", "［填］", "［填］", "［填］", "［填］", "［填］"]],
    [1100] + [982] * 8,
    font_size=8.3,
)

add_heading(doc, "3.1.2 差速转向与PD控制", 3)
add_para(
    doc,
    "循迹控制器根据当前误差及误差变化率计算差速修正量。直线段和弧线段可使用不同的基础速度与参数："
)
add_formula(doc, "Δv(k) = Kₚ·e(k) + K𝒹·[e(k) − e(k−1)] / T")
add_formula(doc, "vL = v₀ + Δv，    vR = v₀ − Δv")
add_para(
    doc,
    "其中v₀为当前路段基础速度，T为控制周期。若弧线存在稳定曲率，可在差速量中加入小幅弯道偏置，但应限制最大修正量，"
    "并通过误差滤波和目标速度斜坡避免轮速突变。"
)

add_heading(doc, "3.1.3 速度闭环", 3)
add_para(
    doc,
    "编码器速度闭环将目标轮速转换为电机PWM。若采用增量式PI，可写为："
)
add_formula(doc, "u(k) = u(k−1) + Kₚᵥ[eᵥ(k)−eᵥ(k−1)] + Kᵢᵥ·eᵥ(k)")
add_para(doc, "输出u(k)需设置限幅、死区补偿和方向控制。报告中应给出采样周期、编码器每圈脉冲数、轮径和最大PWM。")

add_heading(doc, "3.1.4 路段状态机与结束判定", 3)
add_para(
    doc,
    "将环路划分为AB直线、BC弧线、CD直线和DA弧线四个状态。红外组合用于识别路段端部特征，编码器距离用于设置解锁保护，"
    "MPU累计角度用于整圈完成判定。多条件组合可以降低起步或局部黑线造成的误触发。"
)
add_figure_placeholder(
    doc,
    "图5 小车路段状态机",
    "建议节点：IDLE→AB→BC→CD→DA→FINISH；在箭头旁写出实际红外组合、距离阈值、角度阈值和超时保护。",
    height_lines=5,
)
add_table(
    doc,
    ["转移", "主要判据", "保护/备用判据", "实际阈值"],
    [
        ["IDLE→AB", "启动按键按下", "完成传感器校准", "［填写］"],
        ["AB→BC", "［红外组合］", "AB距离解锁/距离保护", "［填写］"],
        ["BC→CD", "重新捕获直线特征", "角度或距离窗口", "［填写］"],
        ["CD→DA", "［红外组合］", "CD距离解锁/距离保护", "［填写］"],
        ["DA→FINISH", "累计角度达到阈值", "启停线或超时保护", "［填写］"],
    ],
    [1250, 2600, 2750, 2357],
    font_size=8.3,
)

add_heading(doc, "3.2 摆杆—钢球系统控制理论", 2)
add_heading(doc, "3.2.1 钢球动力学模型", 3)
add_para(
    doc,
    "将钢球视为在倾斜凹槽中纯滚动的实心球。忽略滚动阻力与槽形高阶影响，沿摆杆方向的加速度近似为："
)
add_formula(doc, "ẍ = (5/7)g·sinθ ≈ (5/7)g·θ")
add_para(
    doc,
    "其中x为钢球相对摆杆中心O的位移，θ为摆杆小角度倾角。该对象本质上具有双积分特性，控制器需限制倾角并抑制超调。"
    "实际凹槽截面、执行机构回差和车体运动会造成模型误差，应通过实验辨识修正。"
)

add_heading(doc, "3.2.2 像素坐标到实际位置", 3)
add_para(doc, "在摆杆刻度上选取多个已知位置，记录钢球圆心像素u，采用线性或分段线性标定：")
add_formula(doc, "x = k·(u − u₀) + b")
add_para(doc, "其中u₀为O点像素坐标。若透视畸变明显，可使用多点二次拟合或单应性变换。标定后应给出均方误差和最大误差。")
add_table(
    doc,
    ["标定位置/cm", "−10", "−5", "0", "+5", "+10", "最大残差/cm"],
    [["像素u", "［填］", "［填］", "［填］", "［填］", "［填］", "［填］"]],
    [1500, 1000, 1000, 1000, 1000, 1000, 2457],
    font_size=8.3,
)

add_heading(doc, "3.2.3 钢球位置闭环", 3)
add_para(
    doc,
    "设定位置为xᵣ，视觉测得位置为x，则位置误差eₓ=xᵣ−x。外环根据位置误差计算目标摆杆角度，执行机构再跟踪目标角度："
)
add_formula(doc, "θ* = Kₚₓeₓ + Kᵢₓ∫eₓdt + K𝒹ₓdeₓ/dt")
add_para(
    doc,
    "应对θ*、角速度及积分项限幅，并对视觉位置进行低通滤波。静止往返测试可采用分段设定点O→+5 cm→−5 cm；"
    "行驶测试则令xᵣ为0或任意指定位置。"
)
add_figure_placeholder(doc, "图6 钢球位置闭环框图", "xᵣ→位置控制器→目标摆杆角度→执行机构/摆杆/钢球→摄像头测量→反馈。")

add_heading(doc, "3.2.4 车体运动扰动补偿", 3)
add_para(
    doc,
    "小车纵向加速度和转弯振动会作为外扰作用于钢球。可利用编码器速度变化或IMU加速度构造前馈补偿θff，最终目标角度为"
    "θcmd=θPID+θff。若实际系统未采用前馈，应删除本段并说明通过速度斜坡、机械减振和位置闭环提高抗扰性。"
)

add_heading(doc, "3.3 参数整定与误差指标", 2)
add_para(
    doc,
    "循迹控制先在低速下确认误差方向，再逐步提高Kₚ直至能够跟线，增加K𝒹抑制蛇形，最后提高基础速度；直线和弧线分别整定。"
    "滚球控制先限制摆杆最大角度，在静止平台上整定位置环，再加入车体运动扰动。"
)
add_formula(doc, "最大绝对误差 Emax = max |x(k) − xᵣ(k)|")
add_formula(doc, "停车偏差 Estop = |s车体基准位置 − sA基准线|")

# Chapter 4
add_heading(doc, "4 电路设计与程序实现", 1)
add_heading(doc, "4.1 硬件电路", 2)
add_para(
    doc,
    "硬件部分应至少给出主控最小系统、电机驱动、编码器、八路红外、MPU6050、视觉通信、摆杆执行机构、按键、OLED/数码管和电源接口。"
    "原理图需标注电压、接口名称、关键上拉/滤波元件以及各模块共地关系。"
)
add_figure_placeholder(doc, "图7 控制系统原理图", "插入清晰原理图；若图过大，可拆分为主控与传感器、电机与执行机构、电源三张。", height_lines=4)
add_table(
    doc,
    ["模块", "实际器件/型号", "主要接口", "供电", "说明"],
    [
        ["主控制器", "［MSPM0G3507/填写］", "［填写］", "［填写］", "控制与任务调度"],
        ["电机驱动", "［TB6612FNG/填写］", "PWM+方向", "［填写］", "四轮驱动"],
        ["循迹模块", "八路红外光电", "8×GPIO", "［填写］", "仅用于黑线循迹"],
        ["惯性测量", "MPU6050", "I²C", "［填写］", "角度与加速度"],
        ["滚球视觉", "［摄像头/处理板］", "［SPI/UART/其他］", "［填写］", "输出钢球坐标"],
        ["摆杆执行", "［填写］", "［PWM/驱动接口］", "［填写］", "控制摆杆倾角"],
        ["显示", "OLED+数码管", "［填写］", "［填写］", "状态与计时"],
    ],
    [1350, 1800, 1650, 1150, 3007],
    font_size=8.0,
)

add_heading(doc, "4.2 软件总体流程", 2)
add_para(
    doc,
    "系统上电后完成GPIO、定时器、显示、编码器、IMU和视觉通信初始化，并执行静态零偏校准。待机时选择测试模式；"
    "按键启动后清零计时、编码器与累计角度，进入对应状态机。控制周期中分别运行循迹/速度控制和滚球位置控制，显示任务时间，"
    "满足结束条件后停止电机与计时并保留最终结果。"
)
add_figure_placeholder(
    doc,
    "图8 主程序流程图",
    "初始化→静态校准→按键选择→清零计时→周期控制→状态判断→停止电机与计时→显示结果。",
    height_lines=5,
)

add_heading(doc, "4.3 小车任务状态机", 2)
add_table(
    doc,
    ["模式/按键", "目标任务", "启动动作", "结束条件", "显示内容"],
    [
        ["KEY1", "一圈并停到A点", "清零计时/角度/里程", "［实际判据］", "运行时间与状态"],
        ["KEY2", "A到B", "清零计时/斜坡起步", "［实际判据］", "AB时间与触发类型"],
        ["KEY3", "慢速一圈", "清零计时/斜坡起步", "［实际判据］", "运行时间与状态"],
        ["滚球任务", "静止/随车定点", "设定xᵣ并启用闭环", "［按实际实现填写］", "设定值/实测值/误差"],
    ],
    [1200, 2500, 2100, 1950, 1207],
    font_size=8.2,
)
add_prompt(doc, "上述按键名称依据当前工程习惯预填；若比赛最终操作方式不同，需按实物改写。")

add_heading(doc, "4.4 循迹与速度控制程序", 2)
add_para(
    doc,
    "循迹任务以［填写控制周期］运行：读取X1～X8，进行有效通道处理和误差滤波，按当前路段选择直线或弧线参数，"
    "计算左右轮目标速度并交给速度环。状态切换只在距离解锁后累计红外特征，避免A点或局部黑线提前触发。"
)
add_figure_placeholder(doc, "图9 循迹控制流程图", "采样→有效性处理→加权误差→滤波→路段参数→PD差速→目标限幅/斜坡→速度环。")

add_heading(doc, "4.5 钢球视觉与控制程序", 2)
add_para(
    doc,
    "视觉端输出钢球圆心及有效标志；主控对丢帧、异常坐标和通信超时进行处理。位置控制任务在新坐标到达时更新误差，"
    "对位置与微分项滤波后计算目标倾角。测试程序记录时间戳、设定位置、实测位置和控制输出，用于绘制误差曲线。"
)
add_figure_placeholder(doc, "图10 钢球检测与控制流程图", "图像采集→ROI→分割/轮廓→坐标标定→滤波→PID→倾角限幅→执行机构。")

add_heading(doc, "4.6 关键参数汇总", 2)
parameter_rows = [
    ["循迹控制周期", "［填写］", "ms", "代码宏/定时器"],
    ["直线基础速度", "［填写］", "%或脉冲/周期", "按KEY模式分列或说明"],
    ["弧线基础速度", "［填写］", "%或脉冲/周期", "按BC、DA分列或说明"],
    ["直线Kₚ、K𝒹", "［填写］", "—", "按KEY模式分列或说明"],
    ["弧线Kₚ、K𝒹", "［填写］", "—", "按BC、DA分列或说明"],
    ["AB/CD距离解锁", "［填写］", "cm", "防止提前触发"],
    ["整圈角度阈值", "［填写］", "°", "实际结束判据"],
    ["滚球控制周期", "［填写］", "ms", "视觉帧率需同时给出"],
    ["滚球Kₚ、Kᵢ、K𝒹", "［填写］", "—", "位置外环"],
    ["最大摆杆角度", "［填写］", "°", "输出限幅"],
]
add_table(doc, ["参数", "数值", "单位", "备注"], parameter_rows, [2700, 1900, 1550, 2807], font_size=8.3)

# Chapter 5
add_heading(doc, "5 测试方案与结果分析", 1)
add_heading(doc, "5.1 测试条件与方法", 2)
add_para(
    doc,
    "测试场地按题目尺寸制作：黑线宽1.8±0.2 cm，AB、CD各1.5 m，BC、DA为半径0.5 m的半圆。"
    "车辆每次从A点相同基准位置出发，电池电压保持在［填写］V。使用［秒表/系统计时］记录时间，直尺测量停车偏差；"
    "钢球位置由摄像头标定坐标逐帧记录，并由录像复核。每项至少测试3次，推荐5次。"
)
add_table(
    doc,
    ["测试项目", "设备/工具", "采样或精度", "说明"],
    [
        ["路线尺寸", "卷尺/直尺", "［填写］", "测试前复核"],
        ["时间", "车载数码管+录像", "0.1 s或更高", "按键启动至任务结束"],
        ["停车偏差", "停车基准虚线+直尺", "1 mm", "测车辆唯一基准点"],
        ["钢球位置", "摄像头+标定程序", "［填写］cm/像素", "计算最大绝对误差"],
        ["电源", "万用表", "［填写］", "记录测试前后电压"],
    ],
    [1800, 2450, 1900, 2807],
    font_size=8.3,
)

add_heading(doc, "5.2 图传显示与录像测试", 2)
add_para(doc, "将接收端置于环形线路外，连续运行各任务，检查画面是否覆盖整根摆杆、有无明显卡顿，并保存录像后随机回放。")
add_table(
    doc,
    ["测试序号", "分辨率/帧率", "是否实时稳定", "是否完整录像", "能否清晰判位", "结论"],
    [
        ["1", "［填写］", "［是/否］", "［是/否］", "［是/否］", "［填写］"],
        ["2", "［填写］", "［是/否］", "［是/否］", "［是/否］", "［填写］"],
        ["3", "［填写］", "［是/否］", "［是/否］", "［是/否］", "［填写］"],
    ],
    [1050, 1550, 1550, 1550, 1750, 1507],
    font_size=8.1,
)
add_figure_placeholder(doc, "图11 图传实时画面与录像回放截图", "截图应能看到整根摆杆、钢球、刻度和录像时间轴。")

add_heading(doc, "5.3 一圈停车测试（任务2）", 2)
add_para(doc, "钢球位置不作要求。小车从A点按键启动，顺时针运行一圈并停止，记录总时间和车体唯一基准点相对A基准线的偏差。")
add_test_table(
    doc,
    ["次数", "总时间/s", "停车偏差/cm", "是否脱线", "计时显示是否保持", "要求"],
    5,
    [750, 1350, 1500, 1250, 2200, 1907],
    "≤20 s；≤2 cm",
    "给出平均时间、最大停车偏差，并分析重复性；若只做3次，可删除多余行。",
)

add_heading(doc, "5.4 静止滚球往返测试（任务3）", 2)
add_para(doc, "小车静止，钢球从O点开始。控制器先令钢球到+5 cm，检测到达后折返到−5 cm并稳定。记录全过程位置曲线。")
add_test_table(
    doc,
    ["次数", "总时间/s", "+5 cm最大误差/cm", "−5 cm最大误差/cm", "最终稳定值/cm", "要求"],
    5,
    [700, 1200, 1950, 1950, 1550, 1607],
    "≤5 s；误差≤1 cm",
    "“运行时间”起止点需统一定义；最大误差从到达指定点后的评价窗口计算时，应说明窗口长度。",
)
add_figure_placeholder(doc, "图12 静止往返位置响应曲线", "横轴时间，纵轴钢球位置；同时画设定值与实测值，并标出±1 cm误差带。")

add_heading(doc, "5.5 A至B动态稳定测试（任务4）", 2)
add_para(doc, "钢球初始置于O点。小车从A点按键启动并通过B点，记录AB时间以及全过程钢球相对O点的最大绝对误差。")
add_test_table(
    doc,
    ["次数", "AB时间/s", "钢球最大误差/cm", "钢球均方根误差/cm", "是否脱线", "要求"],
    5,
    [750, 1300, 1900, 2050, 1200, 1757],
    "≤8 s；误差≤1 cm",
    "报告至少给出最大绝对误差；均方根误差可作为稳定性补充指标。",
)
add_figure_placeholder(doc, "图13 A至B行驶时钢球位置误差曲线", "叠加车辆速度或路段标记，有助于说明启动加速度对钢球的影响。")

add_heading(doc, "5.6 整圈中心点稳定测试（任务5）", 2)
add_para(doc, "钢球设定位置为O点，小车顺时针运行一圈并通过A点。分别标记AB、BC、CD、DA四个路段，统计整圈最大误差。")
add_test_table(
    doc,
    ["次数", "整圈时间/s", "AB最大误差", "BC最大误差", "CD最大误差", "DA最大误差", "要求"],
    5,
    [650, 1150, 1250, 1250, 1250, 1250, 2157],
    "≤30 s；误差≤1 cm",
    "误差单位为cm；若表格过宽，可保留“整圈最大误差”一列，并在曲线中区分路段。",
)
add_figure_placeholder(doc, "图14 整圈中心点稳定位置曲线", "标注AB、BC、CD、DA区间和最大误差出现位置。")

add_heading(doc, "5.7 任意设定位置整圈测试（任务6）", 2)
add_para(
    doc,
    "选择摆杆允许范围内的若干非零设定点进行测试，例如−5 cm、+3 cm、+5 cm。小车从A点运行一圈并通过A点，"
    "记录时间和钢球相对各设定点的最大绝对误差。"
)
add_table(
    doc,
    ["次数", "设定位置/cm", "整圈时间/s", "最大误差/cm", "最终位置/cm", "是否脱线", "要求"],
    [
        ["1", "［填写］", "［填写］", "［填写］", "［填写］", "［是/否］", "≤30 s；≤1 cm"],
        ["2", "［填写］", "［填写］", "［填写］", "［填写］", "［是/否］", "≤30 s；≤1 cm"],
        ["3", "［填写］", "［填写］", "［填写］", "［填写］", "［是/否］", "≤30 s；≤1 cm"],
        ["4", "［填写］", "［填写］", "［填写］", "［填写］", "［是/否］", "≤30 s；≤1 cm"],
        ["5", "［填写］", "［填写］", "［填写］", "［填写］", "［是/否］", "≤30 s；≤1 cm"],
    ],
    [650, 1250, 1200, 1250, 1250, 1050, 2307],
    font_size=8.0,
)
add_figure_placeholder(doc, "图15 不同设定位置的整圈响应曲线", "可在一张图中对比不同设定点，或只展示最具代表性的一组。")

doc.add_page_break()
add_heading(doc, "5.8 测试结果汇总与分析", 2)
add_table(
    doc,
    ["任务", "指标要求", "最差实测结果", "是否达标", "主要限制因素"],
    [
        ["图传", "稳定显示、录像、回放", "［填写］", "［是/否］", "［填写］"],
        ["一圈停车", "≤20 s；≤2 cm", "［填写］", "［是/否］", "［填写］"],
        ["静止往返", "≤5 s；误差≤1 cm", "［填写］", "［是/否］", "［填写］"],
        ["A至B", "≤8 s；误差≤1 cm", "［填写］", "［是/否］", "［填写］"],
        ["整圈中心", "≤30 s；误差≤1 cm", "［填写］", "［是/否］", "［填写］"],
        ["任意位置", "≤30 s；误差≤1 cm", "［填写］", "［是/否］", "［填写］"],
    ],
    [1350, 2250, 2050, 1250, 2057],
    font_size=8.2,
)
add_para(
    doc,
    "结果分析应解释误差出现在哪个阶段及其原因。例如：启动斜坡过陡导致钢球短时偏移，弧线段横向振动使视觉坐标波动，"
    "执行机构回差造成换向滞后，或电池电压下降使速度闭环饱和。随后说明采取的参数或结构改进，并用改进前后数据验证。"
)

# Chapter 6
add_heading(doc, "6 总结", 1)
add_para(
    doc,
    "本文完成了车载平衡滚球运动控制系统的方案设计、理论分析、电路与程序实现，并对题目规定的六项功能进行了测试。"
    "系统通过［概括循迹方法］实现环形线路行驶，通过［概括视觉与滚球控制方法］实现钢球定点与动态稳定。"
    "实测［填写最关键的3～4项结果］，表明系统［满足/基本满足］设计要求。"
)
add_para(
    doc,
    "系统仍存在［填写不足，如视觉延迟、机械回差、强光敏感或高速弯道误差］。后续可通过［填写改进，如提高视觉帧率、"
    "优化摆杆机构、加入扰动前馈或完善路段识别］进一步提高稳定性和重复性。"
)
add_prompt(doc, "总结应以实测结论为主，不再展开新方案；不要写“完全满足全部指标”，除非表5.8确有数据支持。")

# References
add_heading(doc, "参考文献", 1)
refs = [
    "［1］全国大学生电子设计竞赛组委会．H题：车载平衡滚球运动控制系统［Z］．2026．",
    "［2］［主控芯片厂商］．［主控芯片型号］数据手册［EB/OL］．［访问日期］．",
    "［3］TDK InvenSense．MPU-6000 and MPU-6050 Product Specification［EB/OL］．",
    "［4］［电机驱动/摄像头/执行机构的数据手册，按实际补充］．",
    "［5］［控制理论、机器视觉或球杆系统相关参考资料，按实际引用补充］．",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(ref)
    set_run_font(r, east_asia="宋体", size=9.5, color=GRAY if "［" in ref else BLACK)

# Appendices
add_heading(doc, "附录A 主要元器件与成本", 1)
add_table(
    doc,
    ["序号", "名称", "型号/规格", "数量", "用途", "备注"],
    [
        ["1", "主控制器", "［填写］", "1", "整车控制", "［填写］"],
        ["2", "电机驱动", "［填写］", "［填］", "轮式驱动", "［填写］"],
        ["3", "红外循迹模块", "八路光电", "1", "黑线检测", "题目限定"],
        ["4", "惯性传感器", "MPU6050", "1", "角度/加速度", "［填写］"],
        ["5", "摄像头与视觉板", "［填写］", "1套", "钢球检测/图传", "［填写］"],
        ["6", "摆杆执行机构", "［填写］", "1套", "摆杆倾角控制", "［填写］"],
        ["7", "显示装置", "OLED/数码管", "［填］", "状态与计时", "≤2英寸"],
        ["8", "电池与稳压", "［填写］", "1套", "车载供电", "［填写］"],
    ],
    [750, 1600, 1800, 800, 1900, 2107],
    font_size=8.2,
)

add_heading(doc, "附录B 提交前检查清单", 1)
check_items = [
    "摘要中已有主控、循迹、滚球视觉、执行机构和实测结果，且不超过建议篇幅。",
    "全文所有方括号与“填写提示”均已替换或删除。",
    "每张图、每个表都有编号和标题，并在正文中被引用。",
    "所有公式中的符号、单位和参数均在首次出现时解释。",
    "六项任务均有测试方法、至少3组数据、最差结果和达标结论。",
    "报告中的传感器型号、按键功能、阈值和PID参数与最终烧录程序一致。",
    "图传截图能看到整根摆杆、钢球与刻度；录像可完整回放。",
    "停车偏差使用车辆唯一基准点测量，钢球误差使用摄像头标定坐标计算。",
    "删除无实际实现的算法描述，参考文献与正文引用相对应。",
]
for item in check_items:
    add_para(doc, f"□ {item}", first_indent=Cm(0))

add_note_box(
    doc,
    "建议补充材料：",
    "最终报告定稿前准备整车照片、摆杆机构图、原理图、主程序与两个控制流程图、视觉标定截图、六项测试原始数据及位置曲线。"
)

# Prevent table rows from splitting where possible.
for table in doc.tables:
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(str(OUT))
